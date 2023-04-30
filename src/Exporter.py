import docx
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

class IndexExporter:

    def __init__(self):
        pass #Todo eventuell Einstellungen für die Erstellung der Word in die init packen?

    def generate_index(self, words, output_path, grouping):
        index = {}
        for word in words:
            page_numbers = []
            for slide in words[word]:
                page_numbers.append(slide.global_slide_num)
            index[word] = page_numbers
        index = dict(sorted(index.items()))

        # Save the index to a Word document
        self.generate_word_index(index)

    def generate_word_index(self, index):

        doc = docx.Document()

        # Page Settings
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9)

        col_width = Cm(5)
        num_cols = 4

        table = doc.add_table(rows=1, cols=num_cols)
        table.autofit = False
        for col in table.columns:
            col.width = col_width

        # Write the lines to the table
        row = table.rows[0]
        i = 0
        for word, page_numbers in index.items():
            line = f"{word}: {str(page_numbers)}"
            if i >= num_cols:
                i = 0
                table.add_row()
                row = table.rows[-1]
            cell = row.cells[i]
            cell.text = line.strip()
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            cell.paragraphs[0].style.font.name = 'Arial'
            cell.paragraphs[0].style.font.size = Cm(0.3)
            i += 1

        # Save the document
        doc.save('index.docx')


